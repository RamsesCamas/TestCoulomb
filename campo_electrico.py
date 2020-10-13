import math
import numpy as np
from docx2pdf import convert
import docx

def convert_notacion(num):
    num_cientifico = format(num,".2E")
    return num_cientifico

def distancia_o_magnitud(x,y):
    componente = (math.pow(x,2)) + (math.pow(y,2))
    distancia = math.sqrt(componente)
    return distancia

def calcular_angulo(x,y,carga_positiva):
    temp_angle = math.atan2(y,x)
    if carga_positiva and x >0 and y<0:
        temp_angle = math.pi + temp_angle
    if carga_positiva and x<0  and y<0:
        temp_angle = math.pi + temp_angle
    elif carga_positiva and x>0 and y>0:
        temp_angle = math.pi + temp_angle
    return temp_angle


def add_parrafo(doc,parrafo):
    return doc.add_paragraph(parrafo)

def add_next_line(paraObject,line):
    paraObject.add_run(line)

def guardar_doc(doc,nombre):
    file =f'{nombre}.docx'
    doc.save(file)
    return file
    
if __name__ == "__main__":
    K = 9E9
    num_cargas = int(input("Ingrese el número de cargas: "))
    magnitud_cargas = []
    x_cargas = []
    y_cargas = []
    cargas_positivas = []
    for i in range(num_cargas):
        magnitud_carga = float(input("Ingrese la magnitud de la carga: "))
        if magnitud_carga > 0:
            cargas_positivas.append(True)
        else:
            cargas_positivas.append(False)
        magnitud_cargas.append(magnitud_carga)
        x_carga = float(input("Ingrese la coordenada X de la carga: "))
        x_cargas.append(x_carga)
        y_carga = float(input("Ingrese la coordenada Y de la carga: "))
        y_cargas.append(y_carga)
    print("\n")
    x_libre = float(input("Ingrese la coordenada X del punto P: "))
    y_libre = float(input("Ingrese la coordenada Y del punto P: "))
    print("\n")

    doc = docx.Document()
    vector_d_x =[]
    vector_d_y =[]
    distancias = []

    parrafo_cargas = add_parrafo(doc,"Cargas: \n")
    for i in range(len(magnitud_cargas)):
        pos = i +1
        texto_cargas = f'Carga {pos}: {convert_notacion(magnitud_cargas[i])} \t Coordenadas: ({x_cargas[i]},{y_cargas[i]})\n'
        add_next_line(parrafo_cargas,texto_cargas)
    texto_cargas = f'Punto P: \t Coordenadas: ({x_libre},{y_libre})\n'
    add_next_line(parrafo_cargas,texto_cargas)
    


    parrafo_distancia = add_parrafo(doc,"Distancias: \n")
    for i in range(len(magnitud_cargas)):
        pos = i+1
        carga_x = x_cargas[i] - x_libre
        vector_d_x.append(carga_x)
        carga_y =  y_cargas[i] - y_libre
        vector_d_y.append(carga_y)
        d = distancia_o_magnitud(vector_d_x[i], vector_d_y[i])
        texto_distancia = f'Distancia {pos}: {round(d,2)}\n'
        add_next_line(parrafo_distancia,texto_distancia)
        distancias.append(d)

    parrafo_distancia_vector = add_parrafo(doc,"Vector de distancias: \n")
    for i in range(len(vector_d_x)):
        pos = i+1
        Temp_Coord = [vector_d_x[i],vector_d_y[i]]
        text_vec_d = f'Vector de D{pos}: {Temp_Coord}\n'
        add_next_line(parrafo_distancia_vector,text_vec_d)

    angulos_cargas = []
    parrafo_angulos = add_parrafo(doc,"Angulos: \n")
    for i in range(len(cargas_positivas)):
        pos = i+1
        angulo = calcular_angulo(vector_d_x[i],vector_d_y[i],cargas_positivas[i])
        angulos_cargas.append(angulo)
        angulo_grados = math.degrees(angulo)
        texto_angulo = f'Angulo {pos}: {round(angulo_grados,2)}\n'
        add_next_line(parrafo_angulos,texto_angulo)

    parrafo_vector_fuerza = add_parrafo(doc,"Vectores de Fuerza: \n")
    vectores_fuerza = []    
    for i in range(len(distancias)):
        pos = i+1
        temp_magintud = K * magnitud_cargas[i]
        temp_distancia = distancias[i]
        r = math.pow(temp_distancia,2)
        E_x = math.cos(angulos_cargas[i]) * (temp_magintud/r)
        E_y = math.sin(angulos_cargas[i]) * (temp_magintud/r)
        vector_fuerza = np.array([E_x,E_y])
        vectores_fuerza.append(vector_fuerza)
        vector_fuerza_show = f'Vector de Fuerza del Campo {pos}: ' +convert_notacion(E_x) +' , ' + convert_notacion(E_y) +'\n'
        add_next_line(parrafo_vector_fuerza, vector_fuerza_show)
        
    parrafo_fuerza = add_parrafo(doc,"Fuerza del Campo: ")
    for i in range(len(vectores_fuerza)):
        pos = i+1
        F_E = distancia_o_magnitud(vectores_fuerza[i][0],vectores_fuerza[i][1])
        fuerza_show = f'Fuerza de Campo {pos}: '+ convert_notacion(F_E) + '\n'
        add_next_line(parrafo_fuerza,fuerza_show)

    parrafo_resultante = add_parrafo(doc,"Valores Resultantes: \n")
    temp_vector = np.array([0,0])
    for vector in vectores_fuerza:
        temp_vector = np.add(temp_vector,vector)
    Temp_x = convert_notacion(temp_vector[0])
    Temp_y = convert_notacion(temp_vector[1])
    Temp_V = [Temp_x,Temp_y]
    Temp_V = f'\nVector resultante: {Temp_V}\n'
    add_next_line(parrafo_resultante,Temp_V)
    Temp_Mag = distancia_o_magnitud(temp_vector[0],temp_vector[1])
    text_mag = '\nMagnitud resultante: ' + convert_notacion(Temp_Mag) +'\n'
    add_next_line(parrafo_resultante, text_mag)
    Temp_Angle = math.atan2(temp_vector[1],temp_vector[0])
    Temp_Angle = math.degrees(Temp_Angle)
    text_angle = f'Angulo resultante: {round(Temp_Angle,2)}°\n'
    add_next_line(parrafo_resultante, text_angle)
    archivo = guardar_doc(doc,'ExamenCampoElectrico')
    convert(archivo)