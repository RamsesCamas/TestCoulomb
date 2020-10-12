import math
import numpy as np
from docx2pdf import convert
import docx

def convert_notacion(num):
    num_cientifico = format(num,".2E")
    return num_cientifico

def distancia_o_magnitud(x,y):
    componente = (math.pow(x,2)) + (math.pow(y,2))
    distancia = math.sqrt(componente)
    return distancia

def add_parrafo(doc,parrafo):
    return doc.add_paragraph(parrafo)

def add_next_line(paraObject,line):
    paraObject.add_run(line)

def guardar_doc(doc,nombre):
    file =f'{nombre}.docx'
    doc.save(file)
    return file

def convert_pdf(file):
    convert(file)

if __name__ == "__main__":
    K = 9E9
    
    num_cargas = int(input("Ingrese el número de cargas: "))
    magnitud_cargas = []
    x_cargas = []
    y_cargas = []
    for i in range(num_cargas):
        magnitud_carga = float(input("Ingrese la magnitud de la carga: "))
        magnitud_cargas.append(magnitud_carga)
        x_carga = float(input("Ingrese la coordenada X de la carga: "))
        x_cargas.append(x_carga)
        y_carga = float(input("Ingrese la coordenada Y de la carga: "))
        y_cargas.append(y_carga)
    print("\n")
    magnitud_libre = float(input("Ingrese la magnitud de la carga libre: "))
    x_libre = float(input("Ingrese la coordenada X de la carga libre: "))
    y_libre = float(input("Ingrese la coordenada Y de la carga libre: "))
    print("\n")

    doc = docx.Document()
    vector_d_x =[]
    vector_d_y =[]
    distancias = []

    parrafo_cargas = add_parrafo(doc,"Cargas: \n")
    for i in range(len(magnitud_cargas)):
        pos = i +1
        texto_cargas = f'Carga {pos}: {convert_notacion(magnitud_cargas[i])} \t Coordenadas: ({x_cargas[i]},{y_cargas[i]})\n'
        add_next_line(parrafo_cargas,texto_cargas)
    texto_cargas = f'Carga libre: {convert_notacion(magnitud_libre)} \t Coordenadas: ({x_libre},{y_libre})\n'
    add_next_line(parrafo_cargas,texto_cargas)

    parrafo_distancia = add_parrafo(doc,"Distancias: \n")
    for i in range(len(magnitud_cargas)):
        pos = i+1
        carga_x = x_cargas[i] - x_libre
        vector_d_x.append(carga_x)
        carga_y =  y_cargas[i] - y_libre
        vector_d_y.append(carga_y)
        d = distancia_o_magnitud(vector_d_x[i], vector_d_y[i])
        texto_distancia = f'Distancia {pos}: {round(d,2)}\n'
        add_next_line(parrafo_distancia,texto_distancia)
        distancias.append(d)
    parrafo_distancia_vector = add_parrafo(doc,"Vector de distancias: \n")

    for i in range(len(vector_d_x)):
        pos = i+1
        Temp_Coord = [vector_d_x[i],vector_d_y[i]]
        text_vec_d = f'Vector de D{pos}: {Temp_Coord}\n'
        add_next_line(parrafo_distancia_vector,text_vec_d)
    fuerza_cargas = []
    signos_iguales = []

    parrafo_fuerza =add_parrafo(doc,"Fuerza: \n")
    for i in range(len(distancias)):
        pos = i+1
        temp_magintud = magnitud_cargas[i]
        temp_distancia =distancias[i]
        if temp_magintud<0 and magnitud_libre < 0:
            signos_iguales.append(True)
        elif temp_magintud >0 and magnitud_libre >0:
            signos_iguales.append(True)
        else:
            signos_iguales.append(False)
        r = math.pow(temp_distancia,2)
        fuerza = K * ((magnitud_libre*temp_magintud)/r)
        fuerza = abs(fuerza)
        fuerza_cargas.append(fuerza)
        fuerza_show = convert_notacion(fuerza)
        fuerza_show = f'Fuerza {pos}: {fuerza_show}\n'
        add_next_line(parrafo_fuerza,fuerza_show)

    parrafo_angulos = add_parrafo(doc,"Angulos: \n")
    angulos_cargas = []
    angulos_cargas_grados = []
    for i in range(len(distancias)):
        pos = i+1
        angulo = math.atan2(vector_d_y[i],vector_d_x[i])
        if signos_iguales[i]:
            angulo = math.pi + angulo
        angulo_grados = math.degrees(angulo)
        angulos_cargas.append(angulo)
        angulos_cargas_grados.append(angulo_grados) 
        texto_angulo = f'Angulo {pos}: {round(angulo_grados,2)}\n'
        add_next_line(parrafo_angulos,texto_angulo)


    parrafo_vector_fuerza = add_parrafo(doc,"Vectores de Fuerza: \n")
    vectores_fuerza = []    
    new_angles = []
    new_angles_show = []
    for i in range(len(distancias)):
        pos = i+1
        Fx = math.cos(angulos_cargas[i]) * fuerza_cargas[i]
        Fy = math.sin(angulos_cargas[i]) * fuerza_cargas[i]
        Fx_show = convert_notacion(Fx)
        Fy_show = convert_notacion(Fy)
        vector_fuerza_show = [Fx_show,Fy_show]
        vector_fuerza_show = f'Vector de fuerza {pos}: {vector_fuerza_show}\n'
        add_next_line(parrafo_vector_fuerza, vector_fuerza_show)
        vector_fuerza = np.array([Fx,Fy])
        vectores_fuerza.append(vector_fuerza)

    parrafo_resultante = add_parrafo(doc,"Valores Resultantes: \n")
    temp_vector = np.array([0,0])
    for vector in vectores_fuerza:
        temp_vector = np.add(temp_vector,vector)
    Temp_x = convert_notacion(temp_vector[0])
    Temp_y = convert_notacion(temp_vector[1])
    Temp_V = [Temp_x,Temp_y]
    Temp_V = f'\nVector resultante: {Temp_V}\n'
    add_next_line(parrafo_resultante,Temp_V)
    Temp_Mag = distancia_o_magnitud(temp_vector[0],temp_vector[1])
    text_mag = '\nMagnitud resultante: ' + convert_notacion(Temp_Mag) +'\n'
    add_next_line(parrafo_resultante, text_mag)
    Temp_Angle = math.atan2(temp_vector[1],temp_vector[0])
    Temp_Angle = math.degrees(Temp_Angle)
    text_angle = f'Angulo resultante: {round(Temp_Angle,2)}°\n'
    add_next_line(parrafo_resultante, text_angle)
    archivo = guardar_doc(doc,'ExamenCoulomb')
    convert_pdf(archivo)